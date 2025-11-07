"""
Procesador de Documentos - App Carta Manifestación
==================================================

Clase principal para generar cartas de manifestación desde plantilla Word.
"""

import re
from typing import Dict, List, Tuple
from docx import Document


class CartaManifestacionGenerator:
    """Generador de cartas de manifestación a partir de plantilla Word."""

    def __init__(self, template_path: str):
        """
        Inicializa el generador con una plantilla.

        Args:
            template_path: Ruta a la plantilla .docx
        """
        self.template_path = template_path
        self.doc = Document(template_path)
        self.variables = {}
        self.conditionals = {}

    def extract_variables(self) -> Tuple[List[str], List[str]]:
        """
        Extrae todas las variables y condicionales del documento.

        Returns:
            Tuple: (lista de variables, lista de condicionales)
        """
        variables = set()
        conditionals = set()

        # Buscar en párrafos
        for paragraph in self.doc.paragraphs:
            text = paragraph.text

            # Buscar variables {{variable}}
            var_matches = re.findall(r'\{\{([^}]+)\}\}', text)
            for match in var_matches:
                # Limpiar variable
                var_name = match.strip()
                # Caso especial para lista_alto_directores
                if 'lista_alto_directores' in var_name and ':' in var_name:
                    variables.add('lista_alto_directores')
                elif '|' in var_name:  # Manejar filtros como |int
                    var_name = var_name.split('|')[0].strip()
                    if not var_name.startswith('%'):
                        variables.add(var_name)
                elif not var_name.startswith('%'):  # Excluir código Jinja
                    variables.add(var_name)

            # Buscar condicionales {% if variable == 'valor' %}
            cond_matches = re.findall(r'\{%\s*if\s+(\w+)\s*==', text)
            for match in cond_matches:
                conditionals.add(match)

        # Buscar en tablas
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    var_matches = re.findall(r'\{\{([^}]+)\}\}', text)
                    for match in var_matches:
                        var_name = match.strip()
                        if 'lista_alto_directores' in var_name and ':' in var_name:
                            variables.add('lista_alto_directores')
                        elif '|' in var_name:
                            var_name = var_name.split('|')[0].strip()
                            if not var_name.startswith('%'):
                                variables.add(var_name)
                        elif not var_name.startswith('%'):
                            variables.add(var_name)

                    cond_matches = re.findall(r'\{%\s*if\s+(\w+)\s*==', text)
                    for match in cond_matches:
                        conditionals.add(match)

        return sorted(list(variables)), sorted(list(conditionals))

    def _strip_conditional_blocks(self, doc: Document, cond_values: Dict[str, str]) -> None:
        """
        Borra todo el contenido comprendido entre {% if VAR == 'sí' %} ... {% endif %}
        cuando cond_values[VAR] == 'no'.

        Args:
            doc: Documento Word
            cond_values: Diccionario con valores de condiciones
        """
        body_elems = list(doc.element.body)
        inside_remove = False
        inside_keep = False
        trash = []

        for el in body_elems:
            # Obtener texto plano si es párrafo
            txt = ""
            if el.tag.endswith('p'):
                txt = "".join(t.text or "" for t in el.iter() if getattr(t, "text", None)).strip()

            # Apertura del bloque
            m_open = re.match(r"\{% if (\w+)\s*==\s*'sí' %\}", txt)
            if m_open:
                var = m_open.group(1)
                if cond_values.get(var, 'no') == 'sí':
                    inside_keep = True
                else:
                    inside_remove = True
                trash.append(el)
                continue

            # Cierre del bloque
            if re.match(r"\{% endif %\}", txt):
                trash.append(el)
                inside_remove = False
                inside_keep = False
                continue

            # Elementos internos
            if inside_remove:
                trash.append(el)

        # Ejecutar removals
        for el in trash:
            el.getparent().remove(el)

    def process_template(self, variables: Dict[str, str], conditionals: Dict[str, str]) -> Document:
        """
        Procesa la plantilla con las variables proporcionadas.

        Args:
            variables: Diccionario con valores de variables
            conditionals: Diccionario con valores de condicionales

        Returns:
            Document: Documento Word procesado
        """
        new_doc = Document(self.template_path)

        self._strip_conditional_blocks(new_doc, conditionals)

        # Procesar párrafos
        for i, paragraph in enumerate(new_doc.paragraphs):
            original_text = paragraph.text
            if original_text.strip():
                new_text = self._replace_variables(original_text, variables, conditionals)
                if new_text != original_text:
                    # Guardar formato original
                    original_format = self._save_paragraph_format(paragraph)

                    # Limpiar y aplicar nuevo texto
                    paragraph.clear()
                    paragraph.text = new_text

                    # Restaurar formato
                    self._restore_paragraph_format(paragraph, original_format)

        # Procesar tablas
        for table in new_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        if original_text.strip():
                            new_text = self._replace_variables(original_text, variables, conditionals)
                            if new_text != original_text:
                                paragraph.text = new_text

        # Corregir numeración
        self._fix_numbering(new_doc)
        self._remove_underlines(new_doc)
        return new_doc

    def _replace_variables(self, text: str, variables: Dict[str, str], conditionals: Dict[str, str]) -> str:
        """Reemplaza variables y procesa condicionales en un texto."""
        # Primero procesar condicionales
        text = self._process_conditionals(text, conditionals)

        # Manejar caso especial de lista_alto_directores
        lista_pattern = r'\{\{lista_alto_directores:[^}]+\}\}'
        lista_matches = list(re.finditer(lista_pattern, text, re.DOTALL))

        # Reemplazar de atrás hacia adelante
        for match in reversed(lista_matches):
            if 'lista_alto_directores' in variables and variables['lista_alto_directores']:
                text = text[:match.start()] + variables['lista_alto_directores'] + text[match.end():]
            else:
                text = text[:match.start()] + text[match.end():]

        # Reemplazar variables simples
        for var_name, var_value in variables.items():
            if var_name == 'lista_alto_directores':
                continue

            patterns = [
                rf'\{{\{{\s*{re.escape(var_name)}\s*\}}\}}',
                rf'\{{\{{\s*{re.escape(var_name)}\s*\|\s*int\s*\}}\}}',
                rf'\{{\{{\s*{re.escape(var_name)}\s*\|\s*int\s*-\s*1\s*\}}\}}'
            ]

            for pattern in patterns:
                if '|int - 1' in pattern and var_value:
                    try:
                        replacement = str(int(var_value) - 1)
                    except:
                        replacement = var_value
                else:
                    replacement = str(var_value) if var_value else ''

                text = re.sub(pattern, replacement, text)

        # Limpiar marcadores restantes
        text = re.sub(r'\[?\{\{[^}]*\}\}\]?', '', text)
        text = re.sub(r'\[\]\.mark', '', text)
        text = re.sub(r'\.mark', '', text)
        text = re.sub(r'\[\.mark\]', '', text)

        return text

    def _process_conditionals(self, text: str, conditionals: Dict[str, str]) -> str:
        """Procesa bloques condicionales en texto."""
        for cond_var, cond_value in conditionals.items():
            # Patrón para bloques if con mark
            if_pattern = rf'\[\{{% if {cond_var} == \'sí\' %\}}\]\.mark(.*?)\[\{{% endif %\}}\]\.mark'
            if cond_value == 'sí':
                text = re.sub(if_pattern, r'\1', text, flags=re.DOTALL)
            else:
                text = re.sub(if_pattern, '', text, flags=re.DOTALL)

            # Patrón para bloques if sin mark
            if_pattern = rf'\{{% if {cond_var} == \'sí\' %\}}(.*?)\{{% endif %\}}'
            if cond_value == 'sí':
                text = re.sub(if_pattern, r'\1', text, flags=re.DOTALL)
            else:
                text = re.sub(if_pattern, '', text, flags=re.DOTALL)

        # Limpiar marcas de condicionales restantes
        text = re.sub(r'\{%[^%]*%\}', '', text)

        return text

    def _save_paragraph_format(self, paragraph) -> Dict:
        """Guarda el formato de un párrafo."""
        format_info = {
            'alignment': paragraph.alignment,
            'style': paragraph.style.name if paragraph.style else None,
            'runs': []
        }

        for run in paragraph.runs:
            run_format = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
            }
            format_info['runs'].append(run_format)

        return format_info

    def _restore_paragraph_format(self, paragraph, format_info: Dict) -> None:
        """Restaura el formato de un párrafo."""
        if format_info['alignment']:
            paragraph.alignment = format_info['alignment']

        if format_info['style']:
            try:
                paragraph.style = format_info['style']
            except:
                pass

        # Aplicar formato de runs
        if format_info['runs'] and paragraph.runs:
            for i, run in enumerate(paragraph.runs):
                if i < len(format_info['runs']):
                    run_format = format_info['runs'][i]
                    if run_format['bold'] is not None:
                        run.bold = run_format['bold']
                    if run_format['italic'] is not None:
                        run.italic = run_format['italic']
                    if run_format['underline'] is not None:
                        run.underline = run_format['underline']
                    if run_format['font_name']:
                        run.font.name = run_format['font_name']
                    if run_format['font_size']:
                        run.font.size = run_format['font_size']

    def _fix_numbering(self, doc: Document) -> None:
        """Corrige la numeración de los puntos."""
        current_number = 1
        sub_number = 1
        in_sub_list = False

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            # Detectar puntos principales
            main_match = re.match(r'^(\d+)\.\s+(.+)', text)
            if main_match:
                paragraph.text = f"{current_number}. {main_match.group(2)}"
                current_number += 1
                in_sub_list = False

            # Detectar sub-puntos
            sub_match = re.match(r'^[a-z]\.\s+(.+)', text)
            if sub_match:
                if not in_sub_list:
                    sub_number = 1
                    in_sub_list = True

                letter = chr(ord('a') + sub_number - 1)
                paragraph.text = f"{letter}. {sub_match.group(1)}"
                sub_number += 1

    def _remove_underlines(self, doc: Document) -> None:
        """Quita cualquier subrayado de los runs."""
        # Párrafos normales
        for p in doc.paragraphs:
            for run in p.runs:
                run.underline = False

        # Párrafos dentro de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.underline = False
